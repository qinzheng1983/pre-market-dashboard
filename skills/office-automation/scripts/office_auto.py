#!/usr/bin/env python3
"""
Office Automation - 办公自动化工具
支持 Excel、Word、PDF 等常用办公文档处理
"""

import argparse
import json
import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

class ExcelHandler:
    """Excel 处理器"""
    
    def __init__(self):
        self.pd = None
        self.openpyxl = None
        self._init_libs()
        
    def _init_libs(self):
        """初始化库"""
        try:
            import pandas as pd
            self.pd = pd
        except ImportError:
            pass
            
        try:
            import openpyxl
            self.openpyxl = openpyxl
        except ImportError:
            pass
    
    def read_excel(self, filepath: str, sheet_name: str = None) -> Optional[Dict]:
        """读取 Excel 文件"""
        if not self.pd:
            print("❌ 请先安装 pandas: pip install pandas openpyxl")
            return None
            
        try:
            df = self.pd.read_excel(filepath, sheet_name=sheet_name)
            
            # 处理 DataFrame 是 dict 的情况（多个 sheet）
            if isinstance(df, dict):
                sheet_name = list(df.keys())[0]
                df = df[sheet_name]
            
            # 转换为字典
            data = {
                'filename': os.path.basename(filepath),
                'sheet': sheet_name or 'Sheet1',
                'rows': len(df),
                'columns': list(df.columns),
                'data': df.head(10).to_dict('records')  # 只返回前10行
            }
            return data
        except Exception as e:
            print(f"❌ 读取失败: {e}")
            return None
    
    def create_sample(self, output_path: str) -> bool:
        """创建示例 Excel 文件"""
        if not self.pd:
            print("❌ 请先安装 pandas")
            return False
            
        try:
            # 创建示例数据
            data = {
                '日期': ['2026-01-01', '2026-01-02', '2026-01-03'],
                '产品': ['产品A', '产品B', '产品C'],
                '销量': [100, 150, 200],
                '收入': [10000, 15000, 20000]
            }
            df = self.pd.DataFrame(data)
            df.to_excel(output_path, index=False, engine='openpyxl')
            print(f"✅ 已创建: {output_path}")
            return True
        except Exception as e:
            print(f"❌ 创建失败: {e}")
            return False
    
    def excel_to_csv(self, excel_path: str, csv_path: str) -> bool:
        """Excel 转 CSV"""
        if not self.pd:
            print("❌ 请先安装 pandas")
            return False
            
        try:
            df = self.pd.read_excel(excel_path)
            df.to_csv(csv_path, index=False, encoding='utf-8-sig')
            print(f"✅ 已转换: {excel_path} -> {csv_path}")
            return True
        except Exception as e:
            print(f"❌ 转换失败: {e}")
            return False
    
    def csv_to_excel(self, csv_path: str, excel_path: str) -> bool:
        """CSV 转 Excel"""
        if not self.pd:
            print("❌ 请先安装 pandas")
            return False
            
        try:
            df = self.pd.read_csv(csv_path, encoding='utf-8-sig')
            df.to_excel(excel_path, index=False, engine='openpyxl')
            print(f"✅ 已转换: {csv_path} -> {excel_path}")
            return True
        except Exception as e:
            print(f"❌ 转换失败: {e}")
            return False
    
    def merge_excel(self, files: List[str], output_path: str) -> bool:
        """合并多个 Excel 文件"""
        if not self.pd:
            print("❌ 请先安装 pandas")
            return False
            
        try:
            dfs = []
            for f in files:
                df = self.pd.read_excel(f)
                df['来源文件'] = os.path.basename(f)
                dfs.append(df)
            
            merged = self.pd.concat(dfs, ignore_index=True)
            merged.to_excel(output_path, index=False, engine='openpyxl')
            print(f"✅ 已合并 {len(files)} 个文件到: {output_path}")
            return True
        except Exception as e:
            print(f"❌ 合并失败: {e}")
            return False

class WordHandler:
    """Word 处理器"""
    
    def __init__(self):
        self.docx = None
        self._init_lib()
        
    def _init_lib(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            pass
    
    def is_available(self) -> bool:
        """检查是否可用"""
        return self.docx is not None
    
    def create_document(self, output_path: str, title: str = "文档标题", 
                       content: str = "文档内容") -> bool:
        """创建 Word 文档"""
        if not self.docx:
            print("❌ 请先安装 python-docx: pip install python-docx")
            return False
            
        try:
            doc = self.docx.Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            doc.add_paragraph(f"\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.save(output_path)
            print(f"✅ 已创建: {output_path}")
            return True
        except Exception as e:
            print(f"❌ 创建失败: {e}")
            return False
    
    def read_document(self, filepath: str) -> Optional[Dict]:
        """读取 Word 文档"""
        if not self.docx:
            print("❌ 请先安装 python-docx")
            return None
            
        try:
            doc = self.docx.Document(filepath)
            
            # 提取文本
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            return {
                'filename': os.path.basename(filepath),
                'paragraphs': len(doc.paragraphs),
                'content': '\n'.join(full_text[:50])  # 前50行
            }
        except Exception as e:
            print(f"❌ 读取失败: {e}")
            return None

class PDFHandler:
    """PDF 处理器"""
    
    def __init__(self):
        self.fpdf = None
        self._init_lib()
        
    def _init_lib(self):
        try:
            from fpdf import FPDF
            self.fpdf = FPDF
        except ImportError:
            pass
    
    def create_pdf(self, output_path: str, title: str = "PDF 文档",
                  content: str = "文档内容") -> bool:
        """创建 PDF 文档"""
        if not self.fpdf:
            print("❌ 请先安装 fpdf2: pip install fpdf2")
            return False
            
        try:
            pdf = self.fpdf()
            pdf.add_page()
            
            # 添加中文字体支持（简化版）
            try:
                pdf.set_font('Arial', 'B', 16)
            except:
                pdf.set_font('Helvetica', 'B', 16)
            
            pdf.cell(0, 10, title, ln=True, align='C')
            pdf.ln(10)
            
            try:
                pdf.set_font('Arial', '', 12)
            except:
                pdf.set_font('Helvetica', '', 12)
            
            # 简单文本处理（避免中文乱码问题）
            safe_content = content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, safe_content)
            
            pdf.output(output_path)
            print(f"✅ 已创建: {output_path}")
            return True
        except Exception as e:
            print(f"❌ 创建失败: {e}")
            return False

def check_dependencies():
    """检查依赖"""
    print("\n📦 依赖检查")
    print("-" * 40)
    
    deps = [
        ('pandas', 'pandas', 'Excel/CSV 处理'),
        ('openpyxl', 'openpyxl', 'Excel 读写'),
        ('docx', 'python-docx', 'Word 处理'),
        ('fpdf', 'fpdf2', 'PDF 生成')
    ]
    
    for import_name, pkg_name, desc in deps:
        try:
            __import__(import_name)
            print(f"✅ {pkg_name}: {desc}")
        except ImportError:
            print(f"⬜ {pkg_name}: {desc} (未安装)")
    
    print("\n💡 安装命令:")
    print("   pip install pandas openpyxl python-docx fpdf2 --break-system-packages")

def main():
    parser = argparse.ArgumentParser(
        description='Office Automation - 办公自动化工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  # Excel 操作
  office_auto.py excel-read data.xlsx
  office_auto.py excel-create sample.xlsx
  office_auto.py excel-to-csv data.xlsx data.csv
  office_auto.py csv-to-excel data.csv data.xlsx
  
  # Word 操作
  office_auto.py word-create doc.docx "标题" "内容"
  office_auto.py word-read doc.docx
  
  # PDF 操作
  office_auto.py pdf-create doc.pdf "标题" "内容"
  
  # 检查依赖
  office_auto.py check-deps
        """
    )
    
    parser.add_argument('command', nargs='?', help='命令')
    parser.add_argument('args', nargs='*', help='参数')
    
    args = parser.parse_args()
    
    print("=" * 60)
    print("📊 Office Automation - 办公自动化工具")
    print("=" * 60)
    
    if not args.command:
        parser.print_help()
        check_dependencies()
        return
    
    # Excel 操作
    if args.command == 'excel-read':
        if len(args.args) < 1:
            print("❌ 用法: excel-read <文件路径>")
            return
        handler = ExcelHandler()
        result = handler.read_excel(args.args[0])
        if result:
            print(f"\n✅ 读取成功: {result['filename']}")
            print(f"   行数: {result['rows']}")
            print(f"   列名: {', '.join(result['columns'])}")
            print(f"\n前10行数据:")
            for row in result['data']:
                print(f"   {row}")
    
    elif args.command == 'excel-create':
        if len(args.args) < 1:
            print("❌ 用法: excel-create <输出路径>")
            return
        handler = ExcelHandler()
        handler.create_sample(args.args[0])
    
    elif args.command == 'excel-to-csv':
        if len(args.args) < 2:
            print("❌ 用法: excel-to-csv <excel文件> <csv文件>")
            return
        handler = ExcelHandler()
        handler.excel_to_csv(args.args[0], args.args[1])
    
    elif args.command == 'csv-to-excel':
        if len(args.args) < 2:
            print("❌ 用法: csv-to-excel <csv文件> <excel文件>")
            return
        handler = ExcelHandler()
        handler.csv_to_excel(args.args[0], args.args[1])
    
    elif args.command == 'excel-merge':
        if len(args.args) < 2:
            print("❌ 用法: excel-merge <输出文件> <输入文件1> [<输入文件2> ...]")
            return
        handler = ExcelHandler()
        handler.merge_excel(args.args[1:], args.args[0])
    
    # Word 操作
    elif args.command == 'word-create':
        if len(args.args) < 3:
            print("❌ 用法: word-create <输出路径> <标题> <内容>")
            return
        handler = WordHandler()
        handler.create_document(args.args[0], args.args[1], args.args[2])
    
    elif args.command == 'word-read':
        if len(args.args) < 1:
            print("❌ 用法: word-read <文件路径>")
            return
        handler = WordHandler()
        result = handler.read_document(args.args[0])
        if result:
            print(f"\n✅ 读取成功: {result['filename']}")
            print(f"   段落数: {result['paragraphs']}")
            print(f"\n内容预览:\n{result['content'][:500]}...")
    
    # PDF 操作
    elif args.command == 'pdf-create':
        if len(args.args) < 3:
            print("❌ 用法: pdf-create <输出路径> <标题> <内容>")
            return
        handler = PDFHandler()
        handler.create_pdf(args.args[0], args.args[1], args.args[2])
    
    # 依赖检查
    elif args.command == 'check-deps':
        check_dependencies()
    
    else:
        print(f"❌ 未知命令: {args.command}")
        parser.print_help()
    
    print("\n" + "=" * 60)

if __name__ == "__main__":
    main()
